from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "AI-Business-Audit-Premium-Report-Prototype.docx"

FONT = "Aptos"
INK = "18211F"
GREEN = "0B6655"
GREEN_DARK = "08483D"
MINT = "E8F2EF"
PALE = "F5F7F6"
LINE = "CBD5D1"
MUTED = "5F6C68"
WHITE = "FFFFFF"
AMBER = "A46411"
AMBER_PALE = "F7EEDF"
RED = "9C3630"
RED_PALE = "F8E8E6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        data = edges.get(edge)
        if not data:
            continue
        tag = f"w:{edge}"
        node = tc_borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_borders.append(node)
        for key, value in data.items():
            node.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=INK, bold=False, italic=False, all_caps=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    run.font.all_caps = all_caps
    return run


def set_paragraph(paragraph, before=0, after=8, line=1.22, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep
    return paragraph


def add_text(doc, text="", size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8, line=1.22,
             keep=False, all_caps=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    set_paragraph(paragraph, before, after, line, keep)
    set_run(paragraph.add_run(text), size, color, bold, italic, all_caps)
    return paragraph


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 13.5, 3: 11.5}
    colors = {1: GREEN_DARK, 2: GREEN, 3: INK}
    befores = {1: 18, 2: 13, 3: 9}
    afters = {1: 8, 2: 6, 3: 4}
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(p, befores[level], afters[level], 1.1, keep=True)
    set_run(p.add_run(text), sizes[level], colors[level], True)
    return p


def add_kicker(doc, text, color=GREEN, after=8):
    return add_text(doc, text.upper(), 8.5, color, True, False,
                    before=0, after=after, line=1.0, keep=True, all_caps=True)


def add_rule(doc, color=GREEN, size=14, after=12):
    p = doc.add_paragraph()
    set_paragraph(p, 0, after, 1.0)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)
    return p


def add_bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, 0, 5, 1.18)
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_run(p.add_run(text), 10.25, color)
    return p


def add_numbered(doc, number, title, detail):
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [720, 8640])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    left, right = table.rows[0].cells
    set_cell_shading(left, GREEN)
    set_cell_shading(right, PALE)
    for cell in (left, right):
        set_cell_margins(cell, 120, 150, 120, 150)
        set_cell_border(cell, bottom={"val": "single", "sz": 6, "color": LINE})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, 0, 0, 1.0)
    set_run(p.add_run(str(number)), 14, WHITE, True)
    p = right.paragraphs[0]
    set_paragraph(p, 0, 2, 1.15)
    set_run(p.add_run(title), 10.75, INK, True)
    p = right.add_paragraph()
    set_paragraph(p, 0, 0, 1.18)
    set_run(p.add_run(detail), 9.6, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_status_finding(doc, status, title, maturity, impact, fastest_win):
    colors = {
        "RED": (RED, RED_PALE),
        "YELLOW": (AMBER, AMBER_PALE),
        "GREEN": (GREEN, MINT),
    }
    accent, fill = colors[status]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [1540, 7820])
    status_cell, body_cell = table.rows[0].cells
    set_cell_shading(status_cell, accent)
    set_cell_shading(body_cell, fill)
    for cell in (status_cell, body_cell):
        set_cell_margins(cell, 150, 160, 150, 160)
        set_cell_border(cell, bottom={"val": "single", "sz": 6, "color": LINE})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = status_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, 0, 0, 1.0)
    set_run(p.add_run(status), 9, WHITE, True, all_caps=True)
    p = body_cell.paragraphs[0]
    set_paragraph(p, 0, 4, 1.1)
    set_run(p.add_run(title), 11.5, INK, True)
    for label, value in (("Current state", maturity), ("Likely impact", impact), ("Fastest win", fastest_win)):
        p = body_cell.add_paragraph()
        set_paragraph(p, 0, 3, 1.18)
        set_run(p.add_run(f"{label}: "), 9.5, INK, True)
        set_run(p.add_run(value), 9.5, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [9360 // len(metrics)] * len(metrics)
    widths[-1] += 9360 - sum(widths)
    set_table_widths(table, widths)
    for idx, (value, label) in enumerate(metrics):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN_DARK if idx == 0 else PALE)
        set_cell_margins(cell, 150, 130, 150, 130)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            left={"val": "single", "sz": 6, "color": LINE},
            right={"val": "single", "sz": 6, "color": LINE},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, 0, 2, 1.0)
        set_run(p.add_run(value), 19, WHITE if idx == 0 else GREEN_DARK, True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, 0, 0, 1.0)
        set_run(p.add_run(label.upper()), 7.5, WHITE if idx == 0 else MUTED, True, all_caps=True)
    return table


def add_callout(doc, label, headline, detail, fill=MINT, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 200, 220, 200, 220)
    set_cell_border(cell, left={"val": "single", "sz": 24, "color": accent})
    p = cell.paragraphs[0]
    set_paragraph(p, 0, 5, 1.0)
    set_run(p.add_run(label.upper()), 8, accent, True, all_caps=True)
    p = cell.add_paragraph()
    set_paragraph(p, 0, 5, 1.12)
    set_run(p.add_run(headline), 14, INK, True)
    p = cell.add_paragraph()
    set_paragraph(p, 0, 0, 1.2)
    set_run(p.add_run(detail), 10.2, MUTED)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22

    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string({1: GREEN_DARK, 2: GREEN, 3: INK}[level])
        style.font.size = Pt({1: 18, 2: 13.5, 3: 11.5}[level])
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet.font.size = Pt(10.25)
    bullet.paragraph_format.left_indent = Inches(0.38)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)
    bullet.paragraph_format.space_after = Pt(5)
    bullet.paragraph_format.line_spacing = 1.18


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True
    return section


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    set_paragraph(p, 0, 0, 1.0)
    set_run(p.add_run("AI BUSINESS AUDIT"), 8, GREEN, True, all_caps=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.columns[0].width = Inches(4.7)
    table.columns[1].width = Inches(1.8)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, 40, 0, 0, 0)
        set_cell_border(cell, top={"val": "single", "sz": 5, "color": LINE})
    p = left.paragraphs[0]
    set_paragraph(p, 0, 0, 1.0)
    set_run(p.add_run("Confidential | Prepared for BrightSpark Electrical"), 7.5, MUTED)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(p, 0, 0, 1.0)
    set_run(p.add_run("Design prototype | June 2026"), 7.5, MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    section = configure_page(doc)
    add_running_furniture(section)
    core = doc.core_properties
    core.title = "AI Business Audit Premium Report Prototype"
    core.subject = "Revenue and operations readiness report design"
    core.author = "AI Business Audit"
    core.keywords = "business audit, revenue leakage, electrician, automation"

    # Cover
    add_text(doc, "AI BUSINESS AUDIT", 9, GREEN, True, before=14, after=38, line=1.0, all_caps=True)
    add_text(doc, "Revenue & Operations\nReadiness Report", 31, GREEN_DARK, True, after=14, line=0.98)
    add_text(doc, "A practical diagnosis of where leads, follow-up and operational capacity are leaking revenue.", 14, MUTED, after=30, line=1.2)
    add_rule(doc, GREEN, 18, 22)
    add_text(doc, "PREPARED FOR", 8, MUTED, True, after=4, line=1.0, all_caps=True)
    add_text(doc, "BrightSpark Electrical", 19, INK, True, after=3, line=1.05)
    add_text(doc, "Electricians & Small Electrical Contractors", 10.5, MUTED, after=30)
    add_metric_strip(doc, [("6.2", "Readiness score"), ("3", "Priority leaks"), ("30 days", "First action horizon")])
    add_text(doc, "DESIGN PROTOTYPE", 8, AMBER, True, align=WD_ALIGN_PARAGRAPH.RIGHT, before=36, after=2, all_caps=True)
    add_text(doc, "Illustrative content only. Final reports are generated from the customer’s assessment evidence.", 8.5, MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, after=0)

    add_page_break(doc)

    # Executive diagnosis
    add_kicker(doc, "01 | Executive diagnosis")
    add_heading(doc, "The business has demand. The leakage is happening after the enquiry arrives.", 1)
    add_text(
        doc,
        "BrightSpark Electrical appears to have a healthy flow of referral and Google enquiries, but the current response and quote-follow-up process depends too heavily on whoever is available at the time. The result is not a lack of leads; it is inconsistent conversion of the leads already being paid for.",
        after=12,
    )
    add_callout(
        doc,
        "Highest-impact issue",
        "Warm enquiries and issued quotes are not being followed through consistently.",
        "Missed calls are returned manually, quote reminders are not owned by one role, and there is no reliable view of how many opportunities go cold. This creates invisible revenue loss and makes marketing performance look weaker than it may actually be.",
    )
    add_heading(doc, "What matters most now", 2)
    add_numbered(doc, 1, "Recover missed enquiries quickly", "Create an immediate text-back and clear callback ownership for calls that are not answered.")
    add_numbered(doc, 2, "Make quote follow-up systematic", "Use a simple two-step follow-up sequence with visible status and a named owner.")
    add_numbered(doc, 3, "Close the loop after completed work", "Trigger review requests and next-step prompts while customer satisfaction is highest.")
    add_heading(doc, "Assessment confidence", 2)
    add_text(doc, "High. The priority issues are specific, commercially relevant, and can be improved without replacing the entire software stack.", after=4)

    add_page_break(doc)

    # Scores and strengths/gaps
    add_kicker(doc, "02 | Readiness profile")
    add_heading(doc, "A capable business with inconsistent commercial follow-through", 1)
    add_metric_strip(doc, [("6.2", "Overall"), ("5", "Lead response"), ("4", "Follow-up"), ("7", "Process clarity"), ("7", "Customer experience")])
    add_heading(doc, "Where the business is already strong", 2)
    add_bullet(doc, "A clear service offer and strong local reputation create genuine enquiry demand.")
    add_bullet(doc, "The team uses established job-management and accounting tools rather than relying entirely on paper.")
    add_bullet(doc, "Customers receive competent service once a job is booked and allocated.")
    add_heading(doc, "Where revenue is most exposed", 2)
    add_bullet(doc, "After-hours and busy-period calls rely on manual callbacks, with no immediate acknowledgement to the customer.")
    add_bullet(doc, "Quote follow-up varies by staff member and is not measured as a conversion stage.")
    add_bullet(doc, "Review requests and future-work prompts happen inconsistently after completed jobs.")
    add_callout(
        doc,
        "Commercial interpretation",
        "The first investment should improve conversion, not generate more traffic.",
        "Until missed calls and quotes are followed through consistently, increasing ad spend risks feeding more opportunities into the same leakage points.",
        fill=AMBER_PALE,
        accent=AMBER,
    )

    add_page_break(doc)

    # Diagnostic findings
    add_kicker(doc, "03 | Diagnostic findings")
    add_heading(doc, "Evidence-backed issues, ordered by commercial impact", 1)
    add_status_finding(
        doc,
        "RED",
        "Lead response",
        "After-hours and overflow handling is reactive. The owner or office team returns calls when capacity allows.",
        "Urgent customers often call the next electrician quickly. Even a modest number of unrecovered calls can represent significant lost job value.",
        "Send an immediate branded text acknowledgement and assign callback ownership.",
    )
    add_status_finding(
        doc,
        "RED",
        "Quote follow-up",
        "Quotes are sent, but follow-up timing and ownership are inconsistent. Conversion is not visible in one place.",
        "Warm prospects can drift to competitors, while the business cannot distinguish price objections from simple lack of follow-up.",
        "Introduce two scheduled follow-ups and a clear won/lost outcome.",
    )
    add_status_finding(
        doc,
        "YELLOW",
        "Workflow visibility",
        "Core tools are in place, but staff still rely on memory and manual checking to know what needs attention.",
        "The system works while key people are present, but tasks can stall during busy periods or absences.",
        "Create one daily view of uncontacted leads, open quotes and jobs awaiting customer action.",
    )
    add_status_finding(
        doc,
        "YELLOW",
        "Reviews and repeat work",
        "Happy customers are not consistently prompted for a review or relevant future service.",
        "The business misses low-cost reputation growth and repeat-work opportunities after successful jobs.",
        "Trigger a review request after job completion and add one relevant next-step prompt.",
    )
    add_page_break(doc)

    # Action plan
    add_kicker(doc, "04 | 30-day action plan")
    add_heading(doc, "Fix the fastest revenue leaks before adding complexity", 1)
    add_numbered(
        doc, 1, "Days 1–7 | Missed-call recovery",
        "Map the current call path, assign callback ownership, and launch an immediate text acknowledgement for unanswered calls. Track response time and recovered bookings.",
    )
    add_numbered(
        doc, 2, "Days 8–14 | Quote follow-up",
        "Standardise follow-up at two agreed intervals, capture a simple status for every quote, and make open quotes visible to the responsible person.",
    )
    add_numbered(
        doc, 3, "Days 15–21 | Review capture",
        "Send a short review request after completed jobs, with a direct link and one reminder when appropriate.",
    )
    add_numbered(
        doc, 4, "Days 22–30 | Operating visibility",
        "Review the first month of missed calls, quote outcomes and review requests. Keep what works, remove friction, and decide whether deeper automation is justified.",
    )
    add_heading(doc, "Measures to track", 2)
    measures = [
        ("Missed calls acknowledged", "Target: within 1 minute"),
        ("Callback speed", "Target: within 15 minutes during business hours"),
        ("Quotes followed up", "Target: 100% receive the agreed sequence"),
        ("Quote outcomes recorded", "Target: won, lost or pending"),
        ("Review requests sent", "Target: every eligible completed job"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [4680, 4680])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(("Measure", "Initial standard")):
        cell = hdr.cells[idx]
        set_cell_shading(cell, GREEN_DARK)
        set_cell_margins(cell)
        set_cell_border(cell, bottom={"val": "single", "sz": 6, "color": LINE})
        p = cell.paragraphs[0]
        set_paragraph(p, 0, 0, 1.0)
        set_run(p.add_run(text), 9, WHITE, True)
    for measure, standard in measures:
        cells = table.add_row().cells
        for idx, text in enumerate((measure, standard)):
            set_cell_shading(cells[idx], WHITE if len(table.rows) % 2 else PALE)
            set_cell_margins(cells[idx])
            set_cell_border(cells[idx], bottom={"val": "single", "sz": 5, "color": LINE})
            p = cells[idx].paragraphs[0]
            set_paragraph(p, 0, 0, 1.12)
            set_run(p.add_run(text), 9.4, INK if idx == 0 else MUTED, idx == 0)

    add_page_break(doc)

    # Next step / methodology
    add_kicker(doc, "05 | Recommended next step")
    add_heading(doc, "Implement one contained commercial workflow and measure the result", 1)
    add_text(
        doc,
        "The recommended first engagement is a focused follow-up sprint: missed-call recovery, quote follow-up and review capture. It should improve the handling of existing demand without requiring a full CRM replacement or a broad transformation program.",
        after=12,
    )
    add_callout(
        doc,
        "Implementation principle",
        "Automate the hand-off, not the relationship.",
        "Customers should still experience a capable local electrical business. Automation should make acknowledgement, reminders and visibility more reliable while people remain responsible for judgement and customer conversations.",
    )
    add_heading(doc, "How this audit was prepared", 2)
    add_text(
        doc,
        "This report is based on the discovery assessment provided by the business. Findings are limited to the evidence captured during that assessment and are prioritised by likely commercial impact, urgency, implementation effort and speed to visible improvement.",
        after=8,
    )
    add_heading(doc, "Important note", 2)
    add_text(
        doc,
        "The score is a decision aid, not a financial forecast. Any implementation should be confirmed against the business’s actual call volumes, quote data, software configuration and operating constraints.",
        color=MUTED,
        italic=True,
        after=18,
    )
    add_rule(doc, LINE, 6, 12)
    add_text(doc, "AI Business Audit", 13, GREEN_DARK, True, after=3)
    add_text(doc, "Find the revenue leaks. Fix the highest-impact workflow first.", 10, MUTED, after=12)
    add_text(doc, "Prepared for discussion with Hugh.", 9, MUTED, italic=True, after=0)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
